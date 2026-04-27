"""AIcpt — prompt_builder.

Модуль читает описание сети от пользователя (txt/pdf/docx),
склеивает его с pre_prompt.txt, specification_guide.txt и
devices_reference.txt и сохраняет готовый prompt_for_ai.txt
в output/<date_time>/.

Этот модуль намеренно не знает ничего про XML — он только
подготавливает текстовый запрос к нейросети.
"""

from __future__ import annotations

import datetime as _dt
from pathlib import Path


SECTION_SEP = "\n\n" + ("=" * 72) + "\n"


def read_user_input(path: str | Path) -> str:
    """Извлекает текст из .txt / .pdf / .docx. Кодировка txt — utf-8 с заменой."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Файл не найден: {p}")

    ext = p.suffix.lower()
    if ext in (".txt", ".md", ""):
        return p.read_text(encoding="utf-8", errors="replace")

    if ext == ".pdf":
        return _read_pdf(p)

    if ext in (".docx", ".doc"):
        return _read_docx(p)

    raise ValueError(
        f"Неподдерживаемый формат: {ext}. Разрешены: .txt, .pdf, .docx"
    )


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader
        except ImportError as e:
            raise ImportError(
                "Для чтения PDF установи: pip install pypdf"
            ) from e

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
    return "\n".join(parts)


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError(
            "Для чтения DOCX установи: pip install python-docx"
        ) from e

    doc = Document(str(path))
    lines: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(lines)


def _safe_read(path: Path) -> str:
    if not path.exists():
        return f"[предупреждение: файл {path.name} отсутствует]"
    return path.read_text(encoding="utf-8", errors="replace")


def build_prompt(
    user_text: str,
    assets_dir: str | Path,
    topology_description: str | None = None,
) -> str:
    """Склеивает pre_prompt, спецификацию, справочник устройств и текст пользователя.

    `topology_description` — опциональный блок от topology_recogniser_helper,
    содержащий список устройств и связей, распознанных по картинке схемы.
    """
    assets = Path(assets_dir)

    pre_prompt = _safe_read(assets / "pre_prompt.txt")
    spec = _safe_read(assets / "specification_guide.txt")
    devices = _safe_read(assets / "devices_reference.txt")

    blocks = [
        ("# INSTRUCTIONS FOR THE AI", pre_prompt),
        ("# SIMPLIFIED XML SPECIFICATION", spec),
        ("# CISCO PACKET TRACER DEVICE REFERENCE", devices),
    ]
    if topology_description:
        blocks.append((
            "# AUTO-RECOGNISED TOPOLOGY (from diagram in user document, BETA)",
            topology_description,
        ))
    blocks.append(("# USER NETWORK DESCRIPTION (raw)", user_text.strip() or "[empty]"))

    pieces: list[str] = []
    for title, body in blocks:
        pieces.append(title)
        pieces.append("")
        pieces.append(body.strip())
    return SECTION_SEP.join(pieces) + "\n"


def make_output_dir(base: str | Path = "output") -> Path:
    """Создаёт output/YYYY-MM-DD_HH-MM-SS/ и возвращает путь."""
    stamp = _dt.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    out = Path(base) / stamp
    out.mkdir(parents=True, exist_ok=True)
    return out


def run(
    user_file: str | Path,
    assets_dir: str | Path,
    output_base: str | Path = "output",
    auto_topology: bool = False,
    topology_log=None,
) -> tuple[Path, Path]:
    """Полный шаг 1-3: прочитать описание, собрать prompt_for_ai.txt, сохранить.

    auto_topology=True пытается извлечь скриншот схемы из PDF/DOCX и
    добавить распознанное описание устройств/связей в промпт.

    Возвращает (каталог_сессии, путь_до_prompt_for_ai.txt).
    """
    user_text = read_user_input(user_file)

    topology_desc: str | None = None
    if auto_topology:
        ext = Path(user_file).suffix.lower()
        if ext in (".pdf", ".docx", ".doc"):
            try:
                import topology_recogniser_helper as trh
                topology_desc = trh.recognise_from_image(Path(user_file).resolve())
                if topology_log:
                    if topology_desc:
                        n_devs = topology_desc.count("name=")
                        topology_log(
                            f"[+] topology_recogniser: распознано "
                            f"~{n_devs} устройств, добавлено в промпт."
                        )
                    else:
                        topology_log(
                            "[!] topology_recogniser: схему сети не нашёл."
                        )
            except Exception as e:
                if topology_log:
                    topology_log(f"[-] topology_recogniser: {e}")

    prompt = build_prompt(user_text, assets_dir, topology_description=topology_desc)

    session_dir = make_output_dir(output_base)
    prompt_path = session_dir / "prompt_for_ai.txt"
    prompt_path.write_text(prompt, encoding="utf-8")

    (session_dir / "user_input.txt").write_text(user_text, encoding="utf-8")
    if topology_desc:
        (session_dir / "topology_recognised.txt").write_text(
            topology_desc, encoding="utf-8"
        )

    return session_dir, prompt_path


if __name__ == "__main__":
    import argparse

    ap = argparse.ArgumentParser(description="AIcpt: собрать prompt_for_ai.txt")
    ap.add_argument("user_file", help="txt/pdf/docx с описанием сети")
    ap.add_argument(
        "--assets",
        default=str(Path(__file__).parent),
        help="каталог с pre_prompt/spec/devices (по умолчанию рядом со скриптом)",
    )
    ap.add_argument("--out", default="output", help="каталог для output/")
    args = ap.parse_args()

    session, prompt_path = run(args.user_file, args.assets, args.out)
    print(f"[+] Сессия: {session}")
    print(f"[+] Промпт для нейросети: {prompt_path}")
    print("    Скопируй содержимое файла в Claude/ChatGPT и получи упрощённый XML.")
